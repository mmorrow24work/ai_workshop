#!/usr/bin/env python3
"""Convert the AI Workshop markdown guide into a single Word (.docx) document.

v2 changes:
  - Arial 11pt body font
  - Native Word TOC field (updates when opened in Word)
  - All markdown links stripped; Next:/See also: navigation lines removed
  - Code blocks, tables, and headings use keep_together / keep_with_next /
    cantSplit to avoid spanning page breaks
  - Chapter H1s use page_break_before instead of a separate break paragraph
  - README H1 skipped (redundant with cover page)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GUIDE_DIR = Path(__file__).parent / "guide"
RESOURCES_DIR = Path(__file__).parent / "resources"
README = Path(__file__).parent / "README.md"
OUTPUT = Path(__file__).parent / "AI_Workshop_Zero_to_Hero_Guide.docx"

BODY_FONT = "Arial"
BODY_SIZE = 11

FILES_IN_ORDER = [
    README,
    GUIDE_DIR / "01-ai-interfaces.md",
    GUIDE_DIR / "02-llm-history.md",
    GUIDE_DIR / "03-frontier-llms-2026.md",
    GUIDE_DIR / "04-offline-hosting.md",
    GUIDE_DIR / "05-llm-tokens.md",
    GUIDE_DIR / "06-example-packages.md",
    GUIDE_DIR / "07-harness-options.md",
    GUIDE_DIR / "08-security.md",
    GUIDE_DIR / "09-ai-ethics.md",
    GUIDE_DIR / "10-ai-governance.md",
    GUIDE_DIR / "11-ai-news.md",
    GUIDE_DIR / "12-bonus-topics.md",
    RESOURCES_DIR / "youtube-channels.md",
    RESOURCES_DIR / "ai-projects.md",
]


# ─── Link / navigation helpers ───────────────────────────────────────────────

def strip_md_links(text):
    """Replace every [label](url) with just the label text."""
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def is_nav_line(line):
    """Return True for pure navigation lines that should be dropped entirely."""
    s = line.strip()
    # "**Next:** [...](...)" / "*Next:*" / "Next: [...](...)"
    if re.match(r'^\*{0,2}Next\*{0,2}\s*:', s, re.IGNORECASE):
        return True
    # "*See also: [...](...)*"
    if re.match(r'^\*{0,2}See also\*{0,2}\s*:', s, re.IGNORECASE):
        return True
    # A line that is nothing but a markdown link to a .md file
    if re.match(r'^\[.+\]\([^)]+\.md\)\s*$', s):
        return True
    return False


# ─── Paragraph property helpers ──────────────────────────────────────────────

def _remove_and_add(pPr, tag):
    existing = pPr.find(qn(tag))
    if existing is not None:
        pPr.remove(existing)
    el = OxmlElement(tag)
    pPr.append(el)
    return el


def set_keep_together(paragraph):
    """Prevent the paragraph itself from breaking across pages (keepLines)."""
    pPr = paragraph._p.get_or_add_pPr()
    _remove_and_add(pPr, 'w:keepLines')


def set_keep_with_next(paragraph):
    """Keep this paragraph on the same page as the one that follows (keepNext)."""
    pPr = paragraph._p.get_or_add_pPr()
    _remove_and_add(pPr, 'w:keepNext')


def set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    trPr = row._tr.get_or_add_trPr()
    existing = trPr.find(qn('w:cantSplit'))
    if existing is None:
        trPr.append(OxmlElement('w:cantSplit'))


# ─── Inline styling ──────────────────────────────────────────────────────────

def apply_inline_styles(paragraph, text):
    """Add runs to *paragraph* honouring **bold**, *italic*, and `code` markers."""
    text = strip_md_links(text)
    # Order matters: match ** before single *
    pattern = re.compile(r'(\*\*[^*\n]+\*\*|`[^`\n]+`|\*[^*\n]+\*)')
    for part in pattern.split(text):
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)
        else:
            run = paragraph.add_run(part)
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)


# ─── Block elements ──────────────────────────────────────────────────────────

def add_toc(doc):
    """Insert a native Word TOC field that Word populates on open / Ctrl+A F9."""
    heading = doc.add_paragraph("Table of Contents")
    heading.style = "Heading 1"
    for run in heading.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
    set_keep_with_next(heading)

    # Build the TOC field across separate <w:r> elements inside one <w:p>
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pp = p._p

    def append_r(*children):
        r = OxmlElement('w:r')
        for child in children:
            r.append(child)
        pp.append(r)

    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    fc_begin.set(qn('w:dirty'), 'true')   # tells Word to update on open
    append_r(fc_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    append_r(instr)

    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    append_r(fc_sep)

    placeholder = OxmlElement('w:t')
    placeholder.text = (
        "(Right-click this line in Word and choose Update Field "
        "to populate the table of contents.)"
    )
    append_r(placeholder)

    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    append_r(fc_end)

    doc.add_page_break()


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_code_block(doc, code_lines):
    """Shaded monospace block that stays together across page boundaries."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_keep_together(p)  # never split mid-block

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_table(doc, lines):
    """Parse a markdown table and render it as a Word table.

    Separator rows (|---|---|) are skipped.
    All cell text has markdown links stripped.
    Each row has cantSplit set so rows never break mid-row.
    """
    rows = []
    for line in lines:
        if line.startswith('|') and not re.match(r'\|[-| :]+\|', line):
            cells = [strip_md_links(c.strip()) for c in line.strip('|').split('|')]
            rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        set_row_cant_split(row)
        for c_idx in range(col_count):
            cell = row.cells[c_idx]
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
            p = cell.paragraphs[0]
            p.clear()
            apply_inline_styles(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.font.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # breathing room after table


def add_blockquote(doc, text):
    text = strip_md_links(text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_keep_together(p)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '2E75B6')
    pBdr.append(left)
    pPr.append(pBdr)

    apply_inline_styles(p, text)
    for run in p.runs:
        if run.font.name != 'Courier New':
            run.font.italic = True


# ─── File processor ──────────────────────────────────────────────────────────

def process_file(doc, filepath, is_first=False):
    lines = filepath.read_text(encoding='utf-8').splitlines()
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    first_h1_skipped = False   # used to suppress the README's redundant H1

    while i < len(lines):
        line = lines[i]

        # ── Code fence ──────────────────────────────────────────────────
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table accumulation ──────────────────────────────────────────
        if line.startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            add_table(doc, table_lines)
            table_lines = []

        # ── Navigation lines (drop entirely) ────────────────────────────
        if is_nav_line(line):
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.+)', line)
        if m:
            level = len(m.group(1))
            text = strip_md_links(m.group(2).strip())

            if level == 1:
                if is_first and not first_h1_skipped:
                    # README H1 is redundant with the cover page — skip it
                    first_h1_skipped = True
                    i += 1
                    continue
                # Every other H1 starts a new page
                p = doc.add_heading(text, level=1)
                p.paragraph_format.page_break_before = True
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
                set_keep_with_next(p)

            elif level == 2:
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                set_keep_with_next(p)

            elif level == 3:
                p = doc.add_heading(text, level=3)
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                set_keep_with_next(p)

            else:
                p = doc.add_heading(text, level=4)
                for run in p.runs:
                    run.font.name = BODY_FONT
                set_keep_with_next(p)

            i += 1
            continue

        # ── Bullet / numbered list ───────────────────────────────────────
        bm = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)', line)
        if bm:
            indent = len(bm.group(1))
            text = bm.group(3)
            style = 'List Bullet' if indent == 0 else 'List Bullet 2'
            p = doc.add_paragraph(style=style)
            apply_inline_styles(p, text)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # ── Blockquote ──────────────────────────────────────────────────
        if line.startswith('>'):
            add_blockquote(doc, line.lstrip('> ').strip())
            i += 1
            continue

        # ── Empty line ──────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Regular paragraph ───────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        apply_inline_styles(p, line)
        i += 1

    if table_lines:
        add_table(doc, table_lines)


# ─── Document builder ────────────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default body style → Arial 11pt
    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)

    # ── Cover page ──────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(120)
    r = cover.add_run("AI Workshop")
    r.font.name = BODY_FONT
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Zero to Hero Guide")
    r2.font.name = BODY_FONT
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(24)
    r3 = date_p.add_run("June 2026")
    r3.font.name = BODY_FONT
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Native TOC (Word updates this on open or via right-click) ───────
    add_toc(doc)

    # ── Chapter content ─────────────────────────────────────────────────
    for idx, filepath in enumerate(FILES_IN_ORDER):
        if not filepath.exists():
            print(f"  MISSING: {filepath}")
            continue
        print(f"  Processing: {filepath.name}")
        process_file(doc, filepath, is_first=(idx == 0))

    doc.save(OUTPUT)
    print(f"\nSaved: {OUTPUT}")
    print("Note: open in Word and press Ctrl+A then F9 to update the TOC.")


if __name__ == '__main__':
    build_docx()
